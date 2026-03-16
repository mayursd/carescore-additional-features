import json
import os
import re
import time
import io
import streamlit as st
from typing import Optional
from docx import Document

from src.config.prompts import SOAP_PROMPT
from src.services.llm_service import generate_soap_suggestions, llm_call


# ----------------------------
# Utilities
# ----------------------------

def remove_student_name(transcript: str, filename: str) -> str:
    """Remove the student name from the transcript using filename."""
    name_match = re.search(r' - ([A-Za-z]+)\.', filename)
    if name_match:
        student_name = name_match.group(1)
        transcript = re.sub(rf'\b{student_name}\b', '[REDACTED]', transcript, flags=re.IGNORECASE)
    return transcript


def _slice_block(text: str, start_label: str, stop_labels: list[str]) -> str:
    """Return the block after start_label up to the first of stop_labels (or EOF)."""
    start_idx = text.find(start_label)
    if start_idx == -1:
        return ""
    start_idx += len(start_label)
    # Find earliest next stop
    next_positions = [text.find(s, start_idx) for s in stop_labels]
    next_positions = [p for p in next_positions if p != -1]
    end_idx = min(next_positions) if next_positions else len(text)
    return text[start_idx:end_idx].strip()


def _parse_kv_lines(block: str, expected_keys: list[str]) -> dict:
    """
    Parse lines like 'Key: value' into a dict. Only keep expected_keys (case-insensitive match).
    """
    out = {k: "" for k in expected_keys}
    for line in block.splitlines():
        if ":" not in line:
            continue
        k, v = line.split(":", 1)
        key = k.strip()
        val = v.strip()
        # normalize key by underscores/slashes/spaces
        key_norm = key.replace("/", "_").replace(" ", "_")
        for ek in expected_keys:
            if key_norm.lower() == ek.lower():
                out[ek] = val
                break
    return out


def _parse_vitals_from_text(text: str) -> dict:
    """Extract common vital signs from free text into a dict.

    Returns keys: BP, HR, RR, SPO2, Temp, Pain (strings or empty).
    """
    out = {"BP": "", "HR": "", "RR": "", "SPO2": "", "Temp": "", "Pain": ""}
    if not text:
        return out

    t = text.lower()

    # BP patterns: 120/80 or 120 over 80
    import re
    m = re.search(r"(\b(?:bp|blood pressure)[:\s]*)([0-9]{2,3}/[0-9]{2,3})", t, flags=re.IGNORECASE)
    if not m:
        m = re.search(r"([0-9]{2,3})\s*(?:over)\s*([0-9]{2,3})", t)
        if m:
            out["BP"] = f"{m.group(1)}/{m.group(2)}"
    else:
        out["BP"] = m.group(2)

    # HR
    m = re.search(r"\b(?:hr|heart rate)[:\s]*(\d{1,3})\b", t)
    if m:
        out["HR"] = m.group(1)

    # RR
    m = re.search(r"\b(?:rr|respiratory rate)[:\s]*(\d{1,3})\b", t)
    if m:
        out["RR"] = m.group(1)

    # SpO2 / oxygen saturation
    m = re.search(r"\b(?:spo2|o2 sat|oxygen saturation)[:\s]*(\d{1,3}%?)\b", t)
    if m:
        out["SPO2"] = m.group(1)

    # Temp (allow decimal, optional F/C)
    m = re.search(r"\b(?:temp|temperature)[:\s]*(\d{2,3}(?:\.\d)?)\s*(?:f|c)?\b", t)
    if m:
        out["Temp"] = m.group(1)

    # Pain (e.g., 2/10 or score)
    m = re.search(r"\b(?:pain)[:\s]*(\d{1,2}(?:/10)?)\b", t)
    if m:
        out["Pain"] = m.group(1)

    return out


def parse_soap_text_to_dict(text: str) -> dict:
    """
    Parse the single combined SOAP note text from the UI back into a dict that matches
    the template fields expected downstream. Also captures optional AI plan block.
    """
    text = text or ""

    # SUBJECTIVE fields
    hpi = _slice_block(text, "History of Present Illness (HPI):", [
        "\n\nPast Medical History:", "\n\nMedications:", "\n\nFamily History:", "\n\nAllergies:", "\n\nSocial History:"
    ])

    pmhx = _slice_block(text, "Past Medical History:", [
        "\n\nMedications:", "\n\nFamily History:", "\n\nAllergies:", "\n\nSocial History:"
    ])

    meds_subj = _slice_block(text, "Medications:", [
        "\n\nFamily History:", "\n\nAllergies:", "\n\nSocial History:"
    ])

    fhx = _slice_block(text, "Family History:", [
        "\n\nAllergies:", "\n\nSocial History:"
    ])

    allergies = _slice_block(text, "Allergies:", [
        "\n\nSocial History:"
    ])

    shx_block = _slice_block(text, "Social History:", [
        "\n\nReview of Systems:", "\n\n=== OBJECTIVE ===", "\n\n=== ASSESSMENT & PLAN ==="
    ])
    shx_keys = [
        "Tobacco", "ETOH", "Drugs", "Diet", "Exercise",
        "Sexual_activity", "Occupation", "Living_situation", "Safety"
    ]
    shx = _parse_kv_lines(shx_block, shx_keys)

    # ROS
    ros_block = _slice_block(text, "Review of Systems:", [
        "\n\n=== OBJECTIVE ===", "\n\nPhysical Exam:", "\n\n=== ASSESSMENT & PLAN ==="
    ])
    ros_keys = [
        "General", "Eyes", "ENT", "Cardiovascular", "Respiratory", "Gastrointestinal",
        "Genitourinary", "Musculoskeletal", "Neurological", "Psychiatric", "Integument",
        "Endocrine", "Hematopoietic_Lymphatic", "Allergy_Immunology"
    ]
    ros = _parse_kv_lines(ros_block, ros_keys)

    # OBJECTIVE
    obj_block = _slice_block(text, "Physical Exam:", [
        "\n\n=== ASSESSMENT & PLAN ==="
    ])
    obj_keys = [
        "General_Appearance", "HEENT", "Neck", "Cardiovascular", "Pulmonary",
        "GI_Abdomen", "GU", "Musculoskeletal", "Neurological", "Psychiatric", "Integument"
    ]
    obj = _parse_kv_lines(obj_block, obj_keys)

    # ASSESSMENT & PLAN (split into transcript-based vs optional AI block)
    ap_block = _slice_block(text, "=== ASSESSMENT & PLAN ===", [
        # end of document
    ])

    # Detect AI heading and split
    ai_heading_variants = ["AI-Based Assessment & Plan:", "AI-Based Suggestions"]
    ai_start = -1
    ai_label_used = ""
    for label in ai_heading_variants:
        pos = ap_block.find(label)
        if pos != -1:
            ai_start = pos
            ai_label_used = label
            break

    if ai_start != -1:
        ap_main = ap_block[:ai_start]
        ai_plan_text = ap_block[ai_start + len(ai_label_used):].strip()
    else:
        ap_main = ap_block
        ai_plan_text = ""

    ap_keys_all = [
        "Final_diagnosis", "Investigations", "Medications",
        "Education", "Pt_Education", "Follow_Up", "Referrals", "Consults",
        "Disposition", "Other"
    ]
    ap_all = _parse_kv_lines(ap_main, ap_keys_all)
    # Normalize preferred keys
    ap = {
        "Final_diagnosis": ap_all.get("Final_diagnosis", ""),
        "Investigations": ap_all.get("Investigations", ""),
        "Medications": ap_all.get("Medications", ""),
        "Education": ap_all.get("Education", "") or ap_all.get("Pt_Education", ""),
        "Follow_Up": ap_all.get("Follow_Up", ""),
        "Referrals": ap_all.get("Referrals", "") or ap_all.get("Consults", ""),
        "Disposition": ap_all.get("Disposition", ""),
        "Other": ap_all.get("Other", ""),
    }
    if ai_plan_text:
        ap["AI_Plan"] = ai_plan_text

    return {
        "HPI": hpi or "Not documented",
        "PMHx": pmhx or "Not documented",
        "Medications": meds_subj or "Not documented",
        "FHx": fhx or "Not documented",
        "Allergies": allergies or "Not documented",
        "SHx": shx,
        "Review_of_Systems": ros,
        "Objective": obj,
        "Assessment_Plan": ap
    }


# ----------------------------
# LLM extraction
# ----------------------------

def extract_soap_data(transcript, current_texts=None):
    """Extract SOAP note data using Gemini API."""

    current_text_str = ""
    if current_texts:
        current_text_str = "\nCurrent Template Texts:\n" + "\n".join([f"{key}: {value}" for key, value in current_texts.items()])

    soap_prompt = SOAP_PROMPT.format(transcript=transcript, current_text_str=current_text_str)

    selected_model = st.session_state.get('gemini_model', 'gemini-3-pro-preview')

    try:
        response_json = llm_call(
            model=selected_model,
            messages=[
                {"role": "system", "content": "You are a clinician in a medical hospital."},
                {"role": "user", "content": soap_prompt},
            ],
            format="json",
        )
        if not response_json:
            return None

        soap_content = response_json.get("choices", [{}])[0].get("message", {}).get("content", "")
        if not soap_content:
            st.error("❌ Empty Gemini response")
            return None

        if "```" in soap_content:
            soap_content = re.sub(r"^```[a-zA-Z0-9_-]*\s*", "", soap_content.strip())
            soap_content = re.sub(r"\s*```$", "", soap_content.strip())

        parsed_soap = json.loads(soap_content)
        st.write("**Debug SOAP:** ✅ Generated successfully")
        return parsed_soap
    except (json.JSONDecodeError, KeyError) as e:
        st.error(f"❌ Error parsing SOAP data: {e}")
        st.error(f"Raw response: {soap_content}")
        return None
    except Exception as e:
        st.error(f"❌ Gemini error: {str(e)}")
        return None


# ----------------------------
# Public AI plan wrapper (single source of truth)
# ----------------------------

def _normalize_ap_dict(ap_in: Optional[dict]) -> dict:
    ap_in = ap_in or {}
    return {
        "Final_diagnosis": ap_in.get("Final_diagnosis", ""),
        "Investigations": ap_in.get("Investigations", ""),
        "Medications": ap_in.get("Medications", ""),
        "Education": ap_in.get("Education", "") or ap_in.get("Pt_Education", ""),
        "Follow_Up": ap_in.get("Follow_Up", ""),
        "Referrals": ap_in.get("Referrals", "") or ap_in.get("Consults", ""),
        "Disposition": ap_in.get("Disposition", ""),
        "Other": ap_in.get("Other", "")
    }

def generate_ai_assessment_plan(transcript: str, ap_in: dict) -> dict:
    """
    Public API used by both the DOCX and the on-screen editor.
    Delegates to the exact same generator that the DOCX path used before.
    Returns a dict with the same keys as AP.
    """
    ap = _normalize_ap_dict(ap_in)
    case_file_content = getattr(st.session_state, 'case_file_content', '')
    # 🔁 Single source of truth: this is the same logic DOCX used previously
    ai_suggestions = generate_soap_suggestions(transcript, ap, case_file_content) or {}
    # Normalize output to expected keys
    return _normalize_ap_dict(ai_suggestions)


# ----------------------------
# Template population
# ----------------------------

def _read_current_template_texts(doc: Document) -> dict:
    """
    Reads the current visible text in expected template table cells,
    used only to give the LLM some structure hints if needed.
    """
    def safe_cell(tbl_idx, row_idx, col_idx, default=""):
        if tbl_idx < len(doc.tables):
            tbl = doc.tables[tbl_idx]
            if row_idx < len(tbl.rows) and col_idx < len(tbl.rows[row_idx].cells):
                return tbl.rows[row_idx].cells[col_idx].text
        return default

    return {
        "HPI": safe_cell(0, 0, 0, "HPI: (onset, location, duration, character, aggravating/ alleviating, radiation, timing, severity):\n"),
        "PMHx": safe_cell(1, 0, 0, "PMHx: (childhood/ adult illnesses, immunizations, hospitalizations, surgical history):\n"),
        "FHx": safe_cell(1, 0, 1, "FHx: (parents, siblings, children, dx and ages):\n"),
        "SHx": safe_cell(2, 0, 0, "SHx:\nTobacco:  type, number of pack years\nETOH:  type, amount,  frequency,\nDrugs: type, amount, frequency\nDiet:\nExercise:\nSexual activity:\nOccupation:\nLiving situation:\nSafety:\n"),
        "Medications": safe_cell(2, 0, 1, "Medications: (dose, frequency, route, etc, include OTC)\n"),
        "Allergies": safe_cell(2, 1, 1, "Allergies: (specify type of reaction)\n"),
        "Review_of_Systems": safe_cell(3, 0, 0, "Review of Systems:  Include pertinent + and negatives and elaborate where needed.\nGeneral:\nEyes:\nENT:\nCardiovascular:\nRespiratory:\nGastrointestinal:\nGenitourinary:\nMusculoskeletal:\nNeurological:\nPsychiatric:\nIntegument:\nEndocrine:\nHematopoietic/Lymphatic:\nAllergy/Immunology:\n"),
        "Objective": safe_cell(6, 0, 0, "Physical Exam\nGeneral Appearance:\nHEENT:\nNeck:\nCardiovascular:\nPulmonary:\nGI/Abdomen:\nGU:\nMusculoskeletal:\nNeurological:\nPsychiatric:\nIntegument:\n"),
        "Assessment_Plan": safe_cell(8, 0, 0, "Assessment/Plan\nFinal diagnosis or problems(s)\nInvestigations (labs, imaging, etc)\nMedications:  (specific medication, dose, route, frequency)\nConsults:\nDisposition (Admit, d/c, observe, follow up with, when?)\nPt Education:\nOther:\n"),
    }


def populate_soap_template(
    template_file: str,
    transcript: Optional[str] = None,
    ai_suggestions_enabled: bool = True,
    use_ai_assessment_plan: bool = False,
    soap_data: Optional[dict] = None,         # NEW: override from edited text
    soap_text: Optional[str] = None           # optional raw text if you prefer
) -> Optional[bytes]:
    """Populate a SOAP note template with data and return the file bytes."""
    doc = Document(template_file)

    # For LLM prompt scaffolding (optional)
    current_texts = _read_current_template_texts(doc)

    # Final soap_data to use:
    if soap_data and isinstance(soap_data, dict):
        final_soap = soap_data
    elif soap_text:
        final_soap = parse_soap_text_to_dict(soap_text)
    else:
        # Fall back to extracting from transcript
        if not (transcript and transcript.strip()):
            st.error("❌ Cannot populate template: neither soap_data nor transcript provided.")
            return None
        extracted = extract_soap_data(transcript, current_texts)
        if not extracted or "soap_data" not in extracted:
            st.error("Failed to extract SOAP data.")
            return None
        final_soap = extracted["soap_data"]

    # Persist into session for the UI
    st.session_state.soap_data = final_soap

    # Helper to safely update cells
    def update_cell(table_idx, row_idx, cell_idx, new_text, heading):
        if table_idx < len(doc.tables):
            tbl = doc.tables[table_idx]
            if row_idx < len(tbl.rows):
                row = tbl.rows[row_idx]
                if cell_idx < len(row.cells):
                    cell = row.cells[cell_idx]
                    cell.text = f"{heading}{new_text if new_text is not None else ''}"

    def add_heading_to_cell(cell, text, level=2):
        paragraph = cell.add_paragraph()
        run = paragraph.add_run(text)
        paragraph.style = f'Heading {level}'
        return paragraph

    # Safely unpack sections (accept both dict and string for robustness)
    HPI = final_soap.get("HPI", "")
    PMHx = final_soap.get("PMHx", "")
    FHx = final_soap.get("FHx", "")
    Meds_subj = final_soap.get("Medications", "")
    Allergies = final_soap.get("Allergies", "")

    SHx = final_soap.get("SHx", {})
    if not isinstance(SHx, dict):
        # Best effort: parse any "Key: val" lines
        SHx = _parse_kv_lines(str(SHx), [
            "Tobacco", "ETOH", "Drugs", "Diet", "Exercise",
            "Sexual_activity", "Occupation", "Living_situation", "Safety"
        ])
    shx_text = "\n".join([f"{k}: {v}" for k, v in SHx.items()])

    ROS = final_soap.get("Review_of_Systems", {})
    if not isinstance(ROS, dict):
        ROS = _parse_kv_lines(str(ROS), [
            "General", "Eyes", "ENT", "Cardiovascular", "Respiratory", "Gastrointestinal",
            "Genitourinary", "Musculoskeletal", "Neurological", "Psychiatric", "Integument",
            "Endocrine", "Hematopoietic_Lymphatic", "Allergy_Immunology"
        ])
    ros_text = "\n".join([f"{k}: {v}" for k, v in ROS.items()])

    OBJ = final_soap.get("Objective", {})
    if not isinstance(OBJ, dict):
        OBJ = _parse_kv_lines(str(OBJ), [
            "General_Appearance", "HEENT", "Neck", "Cardiovascular", "Pulmonary",
            "GI_Abdomen", "GU", "Musculoskeletal", "Neurological", "Psychiatric", "Integument"
        ])
    objective_text = "\n".join([f"{k}: {v}" for k, v in OBJ.items()])

    AP_in = final_soap.get("Assessment_Plan", {}) or {}
    # Normalize AP keys to support both legacy and preferred names
    AP = {
        "Final_diagnosis": AP_in.get("Final_diagnosis", ""),
        "Investigations": AP_in.get("Investigations", ""),
        "Medications": AP_in.get("Medications", ""),
        "Education": AP_in.get("Education", "") or AP_in.get("Pt_Education", ""),
        "Follow_Up": AP_in.get("Follow_Up", ""),
        "Referrals": AP_in.get("Referrals", "") or AP_in.get("Consults", ""),
        "Disposition": AP_in.get("Disposition", ""),
        "Other": AP_in.get("Other", "")
    }

    # Helper: find a table by header labels (first row). Returns the table or None.
    def _find_table_by_headers(doc_obj, headers):
        for tbl in doc_obj.tables:
            try:
                first_row_texts = [c.text.strip().lower() for c in tbl.rows[0].cells]
            except Exception:
                continue
            match = True
            for h in headers:
                if not any(h.lower() in cell for cell in first_row_texts):
                    match = False
                    break
            if match:
                return tbl
        return None

    # ----------------------------
    # Update Subjective Section
    # ----------------------------
    update_cell(0, 0, 0, HPI, "HPI: (onset, location, duration, character, aggravating/ alleviating, radiation, timing, severity):\n")
    update_cell(1, 0, 0, PMHx, "PMHx: (childhood/ adult illnesses, immunizations, hospitalizations, surgical history):\n")
    update_cell(1, 0, 1, FHx, "FHx: (parents, siblings, children, dx and ages):\n")

    # SHx
    update_cell(2, 0, 0, shx_text, "SHx:\n")

    # Medications and Allergies
    update_cell(2, 0, 1, Meds_subj, "Medications: (dose, frequency, route, etc, include OTC)\n")
    update_cell(2, 1, 1, Allergies, "Allergies: (specify type of reaction)\n")

    # Review of Systems
    update_cell(3, 0, 0, ros_text, "Review of Systems:  Include pertinent + and negatives and elaborate where needed.\n")

    # Objective
    update_cell(6, 0, 0, objective_text, "Physical Exam\n")

    # ----------------------------
    # Vitals: write into a vitals table if present (BP, HR, RR, SPO2, Temp, Pain)
    # ----------------------------
    vitals = final_soap.get("Vitals", {}) or {}
    vitals_table = _find_table_by_headers(doc, ["BP", "HR", "RR", "SPO2", "Temp", "Pain"])
    if vitals_table is not None:
        # Ensure there is a second row to populate; if not, add one (append a row)
        if len(vitals_table.rows) < 2:
            vitals_table.add_row()
        # Map header positions to values
        headers = [c.text.strip() for c in vitals_table.rows[0].cells]
        row_cells = vitals_table.rows[1].cells
        for idx, hdr in enumerate(headers):
            key = hdr.strip().upper()
            val = ""
            if "BP" in key:
                val = vitals.get("BP", "")
            elif "HR" in key:
                val = vitals.get("HR", "")
            elif "RR" in key:
                val = vitals.get("RR", "")
            elif "SPO" in key or "O2" in key:
                val = vitals.get("SPO2", "")
            elif "TEMP" in key:
                val = vitals.get("Temp", "")
            elif "PAIN" in key:
                val = vitals.get("Pain", "")
            # Write into cell (clear first)
            try:
                row_cells[idx].text = str(val or "")
            except Exception:
                pass

    # ----------------------------
    # Assessment/Plan with AI suggestions if enabled
    # ----------------------------
    if ai_suggestions_enabled and use_ai_assessment_plan and (transcript and transcript.strip()):
        # 🔁 Use the shared wrapper so DOCX and UI are identical
        ai_suggestions = generate_ai_assessment_plan(transcript, AP) or {}
    else:
        ai_suggestions = {}

    # Render Assessment/Plan content in Word
    if 8 < len(doc.tables):
        ap_table = doc.tables[8]
        if ap_table.rows and ap_table.rows[0].cells:
            ap_cell = ap_table.cell(0, 0)
            ap_cell.text = ""  # clear existing

            ap_cell.add_paragraph("Assessment/Plan:")

            # Transcript-Based
            add_heading_to_cell(ap_cell, "Transcript-Based Assessment/Plan", level=2)
            transcript_lines = [
                f"Final diagnosis or problems(s): {AP.get('Final_diagnosis', '')}",
                f"Investigations: {AP.get('Investigations', '')}",
                f"Medications: {AP.get('Medications', '')}",
                f"Education: {AP.get('Education', '')}",
                f"Follow-Up: {AP.get('Follow_Up', '')}",
                f"Referrals/Consults: {AP.get('Referrals', '')}",
                f"Disposition: {AP.get('Disposition', '')}",
                f"Other: {AP.get('Other', '')}",
            ]
            for line in transcript_lines:
                ap_cell.add_paragraph(line)

            # AI-Based suggestions (optional)
            if use_ai_assessment_plan and ai_suggestions:
                ap_cell.add_paragraph("")  # spacing
                add_heading_to_cell(ap_cell, "AI-Based Suggestions", level=2)
                ai_lines = [
                    f"Final diagnosis or problems(s): {ai_suggestions.get('Final_diagnosis', '').strip()}",
                    f"Investigations: {ai_suggestions.get('Investigations', '').strip()}",
                    f"Medications: {ai_suggestions.get('Medications', '').strip()}",
                    f"Education: {ai_suggestions.get('Education', '').strip()}",
                    f"Follow-Up: {ai_suggestions.get('Follow_Up', '').strip()}",
                    f"Referrals/Consults: {ai_suggestions.get('Referrals', '').strip()}",
                    f"Disposition: {ai_suggestions.get('Disposition', '').strip()}",
                    f"Other: {ai_suggestions.get('Other', '').strip()}",
                ]
                for line in ai_lines:
                    ap_cell.add_paragraph(line)

    # Signature and Date
    for para in doc.paragraphs:
        if "Student Signature:" in para.text:
            para.text = "Student Signature: "
        if "Date:" in para.text:
            para.text = f"Date: {time.strftime('%Y-%m-%d')}"

    with io.BytesIO() as buffer:
        doc.save(buffer)
        return buffer.getvalue()
